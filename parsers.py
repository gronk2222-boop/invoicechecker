--- glavobsnab_system/parsers.py (原始)


+++ glavobsnab_system/parsers.py (修改后)
"""
parsers.py - Модуль парсинга документов различных форматов
Поддержка: PDF, XLS/XLSX, DOC/DOCX, CSV, TXT
"""

import os
import re
import pandas as pd
from typing import List, Dict, Optional, Tuple
from dataclasses import dataclass


@dataclass
class ParsedItem:
    """Представление позиции номенклатуры"""
    name: str
    quantity: float = None
    unit: str = None
    price_per_unit: float = None
    total_price: float = None
    article: str = None
    gost_standard: str = None
    notes: str = None


@dataclass
class ParsedDocument:
    """Представление parsed документа"""
    supplier_name: str = None
    supplier_address: str = None
    delivery_address: str = None
    items: List[ParsedItem] = None
    total_amount: float = None
    vat_amount: float = None
    delivery_cost: float = None
    free_delivery: bool = False
    currency: str = "RUB"
    raw_text: str = None


def parse_file(file_path: str, file_type: str = None) -> Optional[ParsedDocument]:
    """
    Универсальный парсер файлов
    Определяет тип файла и вызывает соответствующий парсер
    """
    if not os.path.exists(file_path):
        return None

    # Определяем тип файла по расширению если не указан
    if not file_type:
        ext = os.path.splitext(file_path)[1].lower()
        file_type = {
            '.pdf': 'pdf',
            '.xls': 'excel',
            '.xlsx': 'excel',
            '.doc': 'word',
            '.docx': 'word',
            '.csv': 'csv',
            '.txt': 'text'
        }.get(ext, 'text')

    # Вызываем соответствующий парсер
    parsers = {
        'pdf': parse_pdf,
        'excel': parse_excel,
        'word': parse_word,
        'csv': parse_csv,
        'text': parse_text
    }

    parser_func = parsers.get(file_type, parse_text)

    try:
        return parser_func(file_path)
    except Exception as e:
        print(f"Error parsing {file_path}: {e}")
        return None


def parse_pdf(file_path: str) -> Optional[ParsedDocument]:
    """Парсинг PDF файлов с использованием pdfplumber"""
    try:
        import pdfplumber
    except ImportError:
        raise ImportError("pdfplumber not installed. Run: pip install pdfplumber")

    doc = ParsedDocument()
    all_text = []
    tables_data = []

    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            # Извлекаем текст
            text = page.extract_text()
            if text:
                all_text.append(text)

            # Извлекаем таблицы
            tables = page.extract_tables()
            for table in tables:
                tables_data.append(table)

    doc.raw_text = "\n".join(all_text)

    # Парсим информацию из текста и таблиц
    _parse_supplier_info(doc, doc.raw_text)
    _parse_items_from_tables(doc, tables_data)

    # Если items не найдены в таблицах, пробуем извлечь из текста
    if not doc.items:
        _parse_items_from_text(doc, doc.raw_text)

    return doc


def parse_excel(file_path: str) -> Optional[ParsedDocument]:
    """Парсинг Excel файлов (XLS/XLSX)"""
    doc = ParsedDocument()

    # Читаем все листы
    xls = pd.ExcelFile(file_path)
    all_text = []
    all_rows = []

    for sheet_name in xls.sheet_names:
        df = pd.read_excel(xls, sheet_name=sheet_name)
        all_rows.extend(df.values.tolist())
        all_text.append(df.to_string())

    doc.raw_text = "\n".join(all_text)

    # Парсим строки как позиции
    items = []
    for row in all_rows:
        item = _parse_row_to_item(row)
        if item and item.name:
            items.append(item)

    doc.items = items

    # Пытаемся найти информацию о поставщике
    _parse_supplier_info(doc, doc.raw_text)

    # Считаем общую сумму
    if items:
        doc.total_amount = sum(
            item.total_price or (item.quantity * item.price_per_unit if item.quantity and item.price_per_unit else 0)
            for item in items
        )

    return doc


def parse_word(file_path: str) -> Optional[ParsedDocument]:
    """Парсинг Word документов (DOC/DOCX)"""
    try:
        from docx import Document
    except ImportError:
        raise ImportError("python-docx not installed. Run: pip install python-docx")

    doc = ParsedDocument()

    document = Document(file_path)
    all_text = []

    for para in document.paragraphs:
        if para.text.strip():
            all_text.append(para.text)

    # Извлекаем таблицы
    tables_data = []
    for table in document.tables:
        table_rows = []
        for row in table.rows:
            row_cells = [cell.text for cell in row.cells]
            table_rows.append(row_cells)
        tables_data.append(table_rows)

    doc.raw_text = "\n".join(all_text)

    # Парсим информацию
    _parse_supplier_info(doc, doc.raw_text)
    _parse_items_from_tables(doc, tables_data)

    if not doc.items:
        _parse_items_from_text(doc, doc.raw_text)

    return doc


def parse_csv(file_path: str) -> Optional[ParsedDocument]:
    """Парсинг CSV файлов"""
    doc = ParsedDocument()

    try:
        df = pd.read_csv(file_path, encoding='utf-8')
    except UnicodeDecodeError:
        df = pd.read_csv(file_path, encoding='cp1251')

    doc.raw_text = df.to_string()

    # Парсим строки
    items = []
    for _, row in df.iterrows():
        values = row.values.tolist()
        item = _parse_row_to_item(values)
        if item and item.name:
            items.append(item)

    doc.items = items

    if items:
        doc.total_amount = sum(
            item.total_price or (item.quantity * item.price_per_unit if item.quantity and item.price_per_unit else 0)
            for item in items
        )

    return doc


def parse_text(file_path: str) -> Optional[ParsedDocument]:
    """Парсинг текстовых файлов"""
    doc = ParsedDocument()

    encodings = ['utf-8', 'cp1251', 'latin-1']
    content = None

    for encoding in encodings:
        try:
            with open(file_path, 'r', encoding=encoding) as f:
                content = f.read()
            break
        except UnicodeDecodeError:
            continue

    if not content:
        return None

    doc.raw_text = content

    _parse_supplier_info(doc, content)
    _parse_items_from_text(doc, content)

    return doc


# === Вспомогательные функции парсинга ===

def _parse_supplier_info(doc: ParsedDocument, text: str):
    """Извлечение информации о поставщике из текста"""
    if not text:
        return

    lines = text.split('\n')

    # Паттерны для поиска поставщика
    supplier_patterns = [
        r'(?:ООО|АО|ЗАО|ОАО|ИП)\s*["«]?([^"»"]+)["»"]?',
        r'(?:Поставщик|Отправитель|Продавец)[:\s]+(.+)',
    ]

    for pattern in supplier_patterns:
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            doc.supplier_name = match.group(1).strip()
            break

    # Паттерны для адреса
    address_patterns = [
        r'(?:Адрес|Адрес доставки|Доставка)[:\s]+([^\n]+)',
        r'(\d{6},\s*[^\n]+)',  # Почтовый индекс
        r'(г\.|город|ул\.|улица|пр-т|проспект)[^\n]+',
    ]

    for pattern in address_patterns:
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            address = match.group(0).strip()
            if 'доставк' in pattern.lower() or 'доставк' in address.lower():
                doc.delivery_address = address
            elif not doc.supplier_address:
                doc.supplier_address = address
            break

    # Проверка на бесплатную доставку
    if re.search(r'бесплатн.*доставк|доставк.*0\s*руб', text, re.IGNORECASE):
        doc.free_delivery = True

    # Поиск стоимости доставки
    delivery_match = re.search(r'доставк[:а]*\s*(\d+(?:[.,]\d+)?)\s*руб', text, re.IGNORECASE)
    if delivery_match:
        doc.delivery_cost = float(delivery_match.group(1).replace(',', '.'))


def _parse_row_to_item(row: list) -> Optional[ParsedItem]:
    """Преобразование строки таблицы в позицию"""
    if not row or len(row) < 2:
        return None

    # Фильтруем пустые значения
    row = [str(cell).strip() if cell is not None else '' for cell in row]
    row = [cell for cell in row if cell and cell.upper() not in ['ИТОГО', 'ВСЕГО', 'НАИМЕНОВАНИЕ', '№', 'N']]

    if len(row) < 2:
        return None

    item = ParsedItem(name='')

    # Пытаемся определить колонки
    numeric_values = []
    text_values = []

    for val in row:
        # Пробуем распарсить как число
        num_match = re.match(r'^[\d\s]+([.,]\d+)?$', str(val).replace(' ', ''))
        if num_match:
            try:
                numeric_values.append(float(str(val).replace(' ', '').replace(',', '.')))
            except ValueError:
                text_values.append(val)
        else:
            text_values.append(val)

    # Первый текстовый элемент - название
    if text_values:
        item.name = text_values[0].strip()

        # Проверяем на артикул или ГОСТ
        article_match = re.search(r'(арт\.?|артикул|№\s*\d+|GOST|ГОСТ\s*\d+)', item.name, re.IGNORECASE)
        if article_match:
            item.article = article_match.group(0)

    # Числовые значения: количество, цена, сумма
    if len(numeric_values) >= 1:
        item.quantity = numeric_values[0]
    if len(numeric_values) >= 2:
        item.price_per_unit = numeric_values[1]
    if len(numeric_values) >= 3:
        item.total_price = numeric_values[2]

    # Единицы измерения
    unit_patterns = ['шт', 'кг', 'м', 'м2', 'м3', 'л', 'упак', 'короб', 'комплект']
    for val in row:
        for unit in unit_patterns:
            if unit in str(val).lower():
                item.unit = unit
                break

    return item if item.name else None


def _parse_items_from_tables(doc: ParsedDocument, tables_ list):
    """Извлечение позиций из таблиц"""
    items = []

    for table in tables_
        for row in table:
            item = _parse_row_to_item(row)
            if item and item.name and not _is_service_row(item.name):
                items.append(item)

    if items:
        doc.items = items
        doc.total_amount = sum(
            item.total_price or (item.quantity * item.price_per_unit if item.quantity and item.price_per_unit else 0)
            for item in items
        )


def _parse_items_from_text(doc: ParsedDocument, text: str):
    """Извлечение позиций из текста (fallback)"""
    items = []
    lines = text.split('\n')

    # Паттерн для строки с товаром: название, кол-во, цена, сумма
    item_pattern = re.compile(
        r'^(.+?)\s+'
        r'(\d+(?:[.,]\d+)?)\s*'
        r'(шт|кг|м|м2|м3|л|упак)?\s*'
        r'(?:[\d\s]+([.,]\d+)?\s*)?'
        r'(\d+(?:[.,]\d+)?)'
        r'(?:\s*(\d+(?:[.,]\d+)?))?'
    )

    for line in lines:
        match = item_pattern.match(line.strip())
        if match:
            name = match.group(1).strip()
            if not _is_service_row(name):
                item = ParsedItem(
                    name=name,
                    quantity=float(match.group(2).replace(',', '.')),
                    unit=match.group(3),
                    price_per_unit=float(match.group(4).replace(',', '.')) if match.group(4) else None,
                    total_price=float(match.group(5).replace(',', '.')) if match.group(5) else None
                )
                items.append(item)

    if items:
        doc.items = items


def _is_service_row(name: str) -> bool:
    """Проверка является ли строка служебной (не товар)"""
    service_keywords = [
        'итого', 'всего', 'subtotal', 'total', 'ндс', 'налог',
        'доставка', 'transport', 'shipping', 'оплата', 'payment',
        'условия', 'terms', 'реквизиты', 'bank', 'счет', 'invoice',
        '№', 'номер', 'дата', 'date', 'от', 'to', 'для', 'for'
    ]

    name_lower = name.lower()
    return any(keyword in name_lower for keyword in service_keywords)


def extract_items_as_dict(parsed_doc: ParsedDocument) -> List[Dict]:
    """КонвертацияParsedItem в dict для работы с БД"""
    if not parsed_doc or not parsed_doc.items:
        return []

    return [
        {
            'name': item.name,
            'quantity': item.quantity,
            'unit': item.unit,
            'price_per_unit': item.price_per_unit,
            'total_price': item.total_price,
            'article': item.article,
            'gost_standard': item.gost_standard,
            'weight_kg': None,
            'volume_m3': None,
            'characteristics': {},
            'compatibility_info': {}
        }
        for item in parsed_doc.items
    ]
